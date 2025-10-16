"""Document processor orchestrator for contract analysis."""

import logging
import time
import uuid
from pathlib import Path
from typing import Optional, Tuple
import mimetypes

from models.clause import Clause
from models.processed_document import ProcessedDocument
from services.pdf_extractor import PDFExtractor, PDFExtractionError, PasswordProtectedError
from services.ocr_extractor import OCRExtractor, OCRError
from services.clause_segmenter import ClauseSegmenter
from services.google_sheets_service import GoogleSheetsService, GoogleSheetsError
from utils.error_handler import (
    DocumentProcessingError,
    UnsupportedFormatError,
    FileReadError,
    TextExtractionError,
    EmptyDocumentError,
    InputValidator,
    handle_errors
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Main orchestrator for document processing pipeline."""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'],
        'text': ['.txt'],
        'docx': ['.docx']
    }
    
    def __init__(self):
        """Initialize document processor with all required services."""
        self.logger = logging.getLogger(__name__)
        self.pdf_extractor = PDFExtractor()
        self._ocr_extractor = None  # Lazy initialization
        self.clause_segmenter = ClauseSegmenter()
        self._google_sheets_service = None  # Lazy initialization
    
    @property
    def ocr_extractor(self) -> OCRExtractor:
        """Lazy initialization of OCR extractor."""
        if self._ocr_extractor is None:
            try:
                self._ocr_extractor = OCRExtractor(verify_installation=True)
            except OCRError as e:
                self.logger.warning(f"OCR not available: {e}")
                raise
        return self._ocr_extractor
    
    @property
    def google_sheets_service(self) -> GoogleSheetsService:
        """Lazy initialization of Google Sheets service."""
        if self._google_sheets_service is None:
            self._google_sheets_service = GoogleSheetsService()
        return self._google_sheets_service
    
    def process_document(
        self,
        file_path: str,
        file_type: Optional[str] = None,
        use_ocr: bool = False
    ) -> ProcessedDocument:
        """
        Process a document through the complete pipeline.
        
        Args:
            file_path: Path to the document file
            file_type: Optional file type override (pdf, image, text, docx)
            use_ocr: Force OCR extraction even for text-based PDFs
            
        Returns:
            ProcessedDocument with extracted text and segmented clauses
            
        Raises:
            DocumentProcessingError: If processing fails
            UnsupportedFormatError: If file format is not supported
        """
        start_time = time.time()
        file_path = Path(file_path)
        
        # Validate file exists
        if not file_path.exists():
            raise FileReadError(
                f"File not found: {file_path}",
                details={'file_path': str(file_path)}
            )
        
        # Validate file size
        try:
            file_size = file_path.stat().st_size
            InputValidator.validate_file_size(file_size)
        except Exception as e:
            self.logger.error(f"File validation failed: {e}")
            raise
        
        # Detect file type if not provided
        if not file_type:
            file_type = self._detect_file_type(file_path)
        
        self.logger.info(f"Processing document: {file_path.name} (type: {file_type})")
        
        try:
            # Extract text based on file type
            extracted_text, metadata = self._extract_text(
                file_path,
                file_type,
                use_ocr
            )
            
            # Validate extracted text
            if not extracted_text or not extracted_text.strip():
                raise EmptyDocumentError(
                    "No text could be extracted from the document",
                    details={'file_name': file_path.name, 'file_type': file_type}
                )
            
            # Clean the extracted text
            cleaned_text = self._clean_text(extracted_text)
            
            # Validate cleaned text has minimum content
            if len(cleaned_text.strip()) < 50:
                raise EmptyDocumentError(
                    f"Document contains insufficient text ({len(cleaned_text.strip())} characters)",
                    details={'file_name': file_path.name, 'text_length': len(cleaned_text.strip())}
                )
            
            # Segment into clauses
            clauses = self.clause_segmenter.segment(cleaned_text)
            
            # Validate we got some clauses
            if not clauses:
                self.logger.warning(f"No clauses extracted from {file_path.name}")
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            # Create processed document
            document_id = str(uuid.uuid4())
            processed_doc = ProcessedDocument(
                document_id=document_id,
                original_filename=file_path.name,
                extracted_text=cleaned_text,
                clauses=clauses,
                metadata={
                    'file_type': file_type,
                    'file_size': file_size,
                    'extraction_method': metadata.get('extraction_method', 'unknown'),
                    'ocr_confidence': metadata.get('ocr_confidence'),
                    **metadata
                },
                processing_time=processing_time
            )
            
            self.logger.info(
                f"Successfully processed {file_path.name}: "
                f"{processed_doc.num_clauses} clauses, "
                f"{processed_doc.total_words} words, "
                f"{processing_time:.2f}s"
            )
            
            return processed_doc
            
        except (PDFExtractionError, OCRError, UnsupportedFormatError, EmptyDocumentError) as e:
            self.logger.error(f"Document processing failed: {e}")
            raise
        except Exception as e:
            self.logger.exception(f"Unexpected error processing document: {e}")
            raise DocumentProcessingError(
                f"Failed to process document: {e}",
                details={'file_name': file_path.name}
            )
    
    def process_text(self, text: str, filename: str = "pasted_text.txt") -> ProcessedDocument:
        """
        Process raw text input (e.g., pasted contract text).
        
        Args:
            text: Raw contract text
            filename: Optional filename for reference
            
        Returns:
            ProcessedDocument with segmented clauses
            
        Raises:
            EmptyDocumentError: If text is empty or too short
        """
        start_time = time.time()
        
        # Validate text input
        try:
            InputValidator.validate_text_input(text, min_length=100)
        except Exception as e:
            self.logger.error(f"Text validation failed: {e}")
            raise
        
        self.logger.info(f"Processing text input: {len(text)} characters")
        
        try:
            # Clean text
            cleaned_text = self._clean_text(text)
            
            # Validate cleaned text
            if len(cleaned_text.strip()) < 50:
                raise EmptyDocumentError(
                    f"Text contains insufficient content after cleaning ({len(cleaned_text.strip())} characters)",
                    details={'original_length': len(text), 'cleaned_length': len(cleaned_text.strip())}
                )
            
            # Segment into clauses
            clauses = self.clause_segmenter.segment(cleaned_text)
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            # Create processed document
            document_id = str(uuid.uuid4())
            processed_doc = ProcessedDocument(
                document_id=document_id,
                original_filename=filename,
                extracted_text=cleaned_text,
                clauses=clauses,
                metadata={
                    'file_type': 'text',
                    'extraction_method': 'direct_input',
                    'original_length': len(text),
                    'cleaned_length': len(cleaned_text)
                },
                processing_time=processing_time
            )
            
            self.logger.info(
                f"Successfully processed text input: "
                f"{processed_doc.num_clauses} clauses, "
                f"{processing_time:.2f}s"
            )
            
            return processed_doc
            
        except EmptyDocumentError:
            raise
        except Exception as e:
            self.logger.exception(f"Error processing text input: {e}")
            raise DocumentProcessingError(
                f"Failed to process text input: {e}",
                details={'text_length': len(text)}
            )
    
    def process_google_sheet(
        self,
        url: str,
        sheet_name: Optional[str] = None,
        cell_range: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process contract text from Google Sheets.
        
        Args:
            url: Google Sheets URL
            sheet_name: Optional sheet name (defaults to first sheet)
            cell_range: Optional cell range (e.g., 'A1:B10')
            
        Returns:
            ProcessedDocument with segmented clauses
            
        Raises:
            DocumentProcessingError: If processing fails
            GoogleSheetsError: If Google Sheets extraction fails
        """
        start_time = time.time()
        
        # Validate URL
        try:
            InputValidator.validate_google_sheets_url(url)
        except Exception as e:
            self.logger.error(f"Google Sheets URL validation failed: {e}")
            raise
        
        self.logger.info(f"Processing Google Sheets: {url}")
        
        try:
            # Extract text from Google Sheets
            text = self.google_sheets_service.extract_text_from_sheet(
                url,
                sheet_name,
                cell_range
            )
            
            # Validate extracted text
            if not text or not text.strip():
                raise EmptyDocumentError(
                    "No text extracted from Google Sheets",
                    details={'url': url, 'sheet_name': sheet_name, 'cell_range': cell_range}
                )
            
            # Validate minimum text length
            if len(text.strip()) < 100:
                raise EmptyDocumentError(
                    f"Extracted text too short ({len(text.strip())} characters). "
                    "Please ensure the sheet contains contract text (minimum 100 characters).",
                    details={
                        'url': url,
                        'text_length': len(text.strip()),
                        'min_length': 100
                    }
                )
            
            # Clean text
            cleaned_text = self._clean_text(text)
            
            # Segment into clauses
            clauses = self.clause_segmenter.segment(cleaned_text)
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            # Create processed document
            document_id = str(uuid.uuid4())
            
            # Extract spreadsheet ID for filename
            spreadsheet_id, _ = self.google_sheets_service.parse_sheet_url(url)
            filename = f"google_sheet_{spreadsheet_id}.txt"
            
            processed_doc = ProcessedDocument(
                document_id=document_id,
                original_filename=filename,
                extracted_text=cleaned_text,
                clauses=clauses,
                metadata={
                    'file_type': 'google_sheets',
                    'extraction_method': 'google_sheets_api',
                    'sheet_url': url,
                    'sheet_name': sheet_name,
                    'cell_range': cell_range
                },
                processing_time=processing_time
            )
            
            self.logger.info(
                f"Successfully processed Google Sheets: "
                f"{processed_doc.num_clauses} clauses, "
                f"{processing_time:.2f}s"
            )
            
            return processed_doc
            
        except (GoogleSheetsError, EmptyDocumentError) as e:
            self.logger.error(f"Google Sheets processing failed: {e}")
            raise
        except Exception as e:
            self.logger.exception(f"Unexpected error processing Google Sheets: {e}")
            raise DocumentProcessingError(
                f"Failed to process Google Sheets: {e}",
                details={'url': url}
            )
    
    def _extract_text(
        self,
        file_path: Path,
        file_type: str,
        use_ocr: bool
    ) -> Tuple[str, dict]:
        """
        Extract text from file based on type.
        
        Args:
            file_path: Path to file
            file_type: Type of file (pdf, image, text, docx)
            use_ocr: Whether to force OCR extraction
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            TextExtractionError: If text extraction fails
            OCRError: If OCR extraction fails
            UnsupportedFormatError: If file type is not supported
        """
        metadata = {}
        
        if file_type == 'pdf':
            if use_ocr:
                # Force OCR extraction
                try:
                    text, confidence = self.ocr_extractor.extract_text_from_pdf(str(file_path))
                    metadata['extraction_method'] = 'ocr'
                    metadata['ocr_confidence'] = confidence
                except OCRError as e:
                    self.logger.error(f"OCR extraction failed: {e}")
                    raise OCRError(
                        f"OCR extraction failed. Tesseract may not be installed or the image quality is poor: {e}",
                        details={'file_name': file_path.name}
                    )
            else:
                # Try regular extraction first
                try:
                    text = self.pdf_extractor.extract_text(str(file_path))
                    metadata['extraction_method'] = 'text_extraction'
                except PDFExtractionError as e:
                    # Fallback to OCR
                    self.logger.info(f"Text extraction failed, trying OCR fallback: {e}")
                    try:
                        text, confidence = self.ocr_extractor.extract_text_from_pdf(str(file_path))
                        metadata['extraction_method'] = 'ocr_fallback'
                        metadata['ocr_confidence'] = confidence
                        self.logger.info(f"OCR fallback successful with confidence: {confidence:.2f}")
                    except OCRError as ocr_e:
                        self.logger.error(f"Both extraction methods failed")
                        raise TextExtractionError(
                            f"Unable to extract text from PDF. Both text extraction and OCR failed.",
                            details={
                                'file_name': file_path.name,
                                'text_extraction_error': str(e),
                                'ocr_error': str(ocr_e)
                            }
                        )
        
        elif file_type == 'image':
            try:
                text, confidence = self.ocr_extractor.extract_text_from_image(str(file_path))
                metadata['extraction_method'] = 'ocr'
                metadata['ocr_confidence'] = confidence
                
                # Warn if confidence is low
                if confidence < 0.7:
                    self.logger.warning(
                        f"Low OCR confidence ({confidence:.2f}) for {file_path.name}. "
                        "Results may be inaccurate."
                    )
            except OCRError as e:
                self.logger.error(f"OCR extraction from image failed: {e}")
                raise OCRError(
                    f"Unable to extract text from image. Ensure the image is clear and contains readable text: {e}",
                    details={'file_name': file_path.name}
                )
        
        elif file_type == 'text':
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                metadata['extraction_method'] = 'text_file'
            except Exception as e:
                self.logger.error(f"Failed to read text file: {e}")
                raise FileReadError(
                    f"Unable to read text file: {e}",
                    details={'file_name': file_path.name}
                )
        
        elif file_type == 'docx':
            try:
                text = self._extract_from_docx(file_path)
                metadata['extraction_method'] = 'docx'
            except Exception as e:
                self.logger.error(f"Failed to extract from DOCX: {e}")
                raise TextExtractionError(
                    f"Unable to extract text from DOCX file: {e}",
                    details={'file_name': file_path.name}
                )
        
        else:
            raise UnsupportedFormatError(
                f"Unsupported file type: {file_type}",
                details={
                    'file_type': file_type,
                    'supported_formats': ', '.join(self._get_all_extensions())
                }
            )
        
        return text, metadata
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return '\n\n'.join(paragraphs)
            
        except ImportError:
            raise DocumentProcessingError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract text from DOCX: {e}")
    
    def _detect_file_type(self, file_path: Path) -> str:
        """
        Detect file type from extension and MIME type.
        
        Args:
            file_path: Path to file
            
        Returns:
            File type (pdf, image, text, docx)
        """
        extension = file_path.suffix.lower()
        
        # Check against supported formats
        for file_type, extensions in self.SUPPORTED_FORMATS.items():
            if extension in extensions:
                return file_type
        
        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            if mime_type.startswith('image/'):
                return 'image'
            elif mime_type == 'application/pdf':
                return 'pdf'
            elif mime_type == 'text/plain':
                return 'text'
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return 'docx'
        
        raise UnsupportedFormatError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: {', '.join(self._get_all_extensions())}"
        )
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace while preserving paragraph structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Normalize spaces within line
                line = ' '.join(line.split())
                cleaned_lines.append(line)
        
        # Join with single newlines
        cleaned = '\n'.join(cleaned_lines)
        
        # Remove common artifacts
        cleaned = cleaned.replace('\x00', '')  # Null bytes
        cleaned = cleaned.replace('\ufffd', '')  # Replacement character
        cleaned = cleaned.replace('\r', '')  # Carriage returns
        
        return cleaned
    
    def _get_all_extensions(self) -> list:
        """Get list of all supported file extensions."""
        extensions = []
        for exts in self.SUPPORTED_FORMATS.values():
            extensions.extend(exts)
        return extensions
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if format is supported
        """
        try:
            self._detect_file_type(Path(file_path))
            return True
        except UnsupportedFormatError:
            return False

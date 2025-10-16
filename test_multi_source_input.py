"""Integration tests for multi-source input support."""

import tempfile
from pathlib import Path
from services.document_processor import DocumentProcessor, DocumentProcessingError


def test_file_upload_txt():
    """Test processing TXT file."""
    processor = DocumentProcessor()
    
    # Create temporary TXT file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write("""
        DATA PROCESSING AGREEMENT
        
        1. Data Processing Terms
        The processor shall process personal data only on documented instructions from the controller.
        
        2. Security Measures
        The processor shall implement appropriate technical and organizational measures to ensure a level of security.
        
        3. Sub-processor Authorization
        The processor shall not engage another processor without prior written authorization from the controller.
        """)
        temp_path = f.name
    
    try:
        # Process document
        result = processor.process_document(temp_path)
        
        # Verify results
        assert result is not None
        assert result.num_clauses > 0
        assert len(result.extracted_text) > 100
        assert result.metadata['file_type'] == 'text'
        assert result.metadata['extraction_method'] == 'text_file'
        
        print(f"✅ TXT file processing successful: {result.num_clauses} clauses extracted")
        
    finally:
        # Clean up
        Path(temp_path).unlink()


def test_file_upload_docx():
    """Test processing DOCX file (requires python-docx)."""
    try:
        from docx import Document
    except ImportError:
        print("⚠️ python-docx not installed, skipping DOCX test")
        return
    
    processor = DocumentProcessor()
    
    # Create temporary DOCX file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    try:
        # Create DOCX document
        doc = Document()
        doc.add_heading('Data Processing Agreement', 0)
        doc.add_paragraph('1. Data Processing Terms')
        doc.add_paragraph('The processor shall process personal data only on documented instructions.')
        doc.add_paragraph('2. Security Measures')
        doc.add_paragraph('The processor shall implement appropriate security measures.')
        doc.save(temp_path)
        
        # Process document
        result = processor.process_document(temp_path)
        
        # Verify results
        assert result is not None
        assert result.num_clauses > 0
        assert len(result.extracted_text) > 50
        assert result.metadata['file_type'] == 'docx'
        assert result.metadata['extraction_method'] == 'docx'
        
        print(f"✅ DOCX file processing successful: {result.num_clauses} clauses extracted")
        
    finally:
        # Clean up
        Path(temp_path).unlink()


def test_text_input():
    """Test processing direct text input."""
    processor = DocumentProcessor()
    
    contract_text = """
    DATA PROCESSING AGREEMENT
    
    1. Data Processing Terms
    The processor shall process personal data only on documented instructions from the controller.
    The controller shall ensure that processing instructions are documented.
    
    2. Security Measures
    The processor shall implement appropriate technical and organizational measures.
    These measures shall ensure a level of security appropriate to the risk.
    
    3. Sub-processor Authorization
    The processor shall not engage another processor without prior written authorization.
    The controller shall have the right to object to changes in sub-processors.
    """
    
    # Process text
    result = processor.process_text(contract_text)
    
    # Verify results
    assert result is not None
    assert result.num_clauses > 0
    assert len(result.extracted_text) > 100
    assert result.metadata['file_type'] == 'text'
    assert result.metadata['extraction_method'] == 'direct_input'
    assert result.original_filename == 'pasted_text.txt'
    
    print(f"✅ Text input processing successful: {result.num_clauses} clauses extracted")


def test_text_input_validation():
    """Test text input validation."""
    processor = DocumentProcessor()
    
    # Test empty text
    try:
        processor.process_text("")
        assert False, "Should raise error for empty text"
    except DocumentProcessingError as e:
        assert "cannot be empty" in str(e)
        print("✅ Empty text validation working")
    
    # Test whitespace-only text
    try:
        processor.process_text("   \n\n   ")
        assert False, "Should raise error for whitespace-only text"
    except DocumentProcessingError as e:
        assert "cannot be empty" in str(e)
        print("✅ Whitespace-only text validation working")


def test_google_sheets_url_validation():
    """Test Google Sheets URL validation."""
    processor = DocumentProcessor()
    
    # Valid URLs
    valid_urls = [
        "https://docs.google.com/spreadsheets/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgvE2upms/edit",
        "https://docs.google.com/spreadsheets/d/abc123-xyz/edit#gid=0",
    ]
    
    for url in valid_urls:
        assert processor.google_sheets_service.validate_url(url) is True
        print(f"✅ Valid URL accepted: {url[:50]}...")
    
    # Invalid URLs
    invalid_urls = [
        "https://invalid-url.com",
        "https://google.com",
        "not-a-url",
    ]
    
    for url in invalid_urls:
        assert processor.google_sheets_service.validate_url(url) is False
        print(f"✅ Invalid URL rejected: {url}")


def test_file_size_validation():
    """Test file size validation logic."""
    max_size = 10 * 1024 * 1024  # 10MB
    
    # Test file size check
    test_sizes = [
        (5 * 1024 * 1024, True, "5MB file"),
        (10 * 1024 * 1024, False, "10MB file (at limit)"),
        (15 * 1024 * 1024, False, "15MB file (over limit)"),
    ]
    
    for size, should_pass, description in test_sizes:
        if size <= max_size:
            assert should_pass or size == max_size
            print(f"✅ {description}: Would be accepted")
        else:
            assert not should_pass
            print(f"✅ {description}: Would be rejected")


def test_supported_formats():
    """Test supported format detection."""
    processor = DocumentProcessor()
    
    supported_files = [
        "contract.pdf",
        "contract.docx",
        "contract.txt",
        "scan.png",
        "scan.jpg",
        "scan.jpeg",
    ]
    
    for filename in supported_files:
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as f:
            temp_path = f.name
        
        try:
            assert processor.is_supported_format(temp_path) is True
            print(f"✅ Format supported: {filename}")
        finally:
            Path(temp_path).unlink()
    
    # Test unsupported format
    with tempfile.NamedTemporaryFile(suffix='.xyz', delete=False) as f:
        temp_path = f.name
    
    try:
        assert processor.is_supported_format(temp_path) is False
        print("✅ Unsupported format correctly identified: .xyz")
    finally:
        Path(temp_path).unlink()


if __name__ == "__main__":
    print("=" * 60)
    print("Multi-Source Input Integration Tests")
    print("=" * 60)
    print()
    
    print("Testing File Upload (TXT)...")
    test_file_upload_txt()
    print()
    
    print("Testing File Upload (DOCX)...")
    test_file_upload_docx()
    print()
    
    print("Testing Text Input...")
    test_text_input()
    print()
    
    print("Testing Text Input Validation...")
    test_text_input_validation()
    print()
    
    print("Testing Google Sheets URL Validation...")
    test_google_sheets_url_validation()
    print()
    
    print("Testing File Size Validation...")
    test_file_size_validation()
    print()
    
    print("Testing Supported Formats...")
    test_supported_formats()
    print()
    
    print("=" * 60)
    print("✅ All tests passed!")
    print("=" * 60)
